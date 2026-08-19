"""Safe storage and text extraction for uploaded documents."""

from __future__ import annotations

import re
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from uuid import UUID, uuid4

from docx import Document
from pypdf import PdfReader

from app.core.exceptions import (
    DocumentTooLargeError,
    EmptyDocumentError,
    InvalidDocumentError,
    UnsupportedDocumentTypeError,
)

SUPPORTED_DOCUMENT_SUFFIXES = frozenset({".pdf", ".docx"})


@dataclass(frozen=True)
class StoredDocument:
    """Stored upload metadata and extracted plain text."""

    original_name: str
    path: Path
    text: str


class DocumentFileService:
    """Validate, store, and extract text from PDF or DOCX uploads."""

    def __init__(self, upload_dir: Path, max_size_mb: int) -> None:
        self._upload_dir = upload_dir
        self._max_size_bytes = max_size_mb * 1024 * 1024

    def process(self, user_id: UUID, file_name: str, content: bytes) -> StoredDocument:
        """Validate an upload, extract its text, and persist it."""
        safe_name = self._safe_file_name(file_name)
        suffix = Path(safe_name).suffix.lower()
        if suffix not in SUPPORTED_DOCUMENT_SUFFIXES:
            raise UnsupportedDocumentTypeError("Only PDF and DOCX files are supported")
        if not content:
            raise EmptyDocumentError("The uploaded document is empty")
        if len(content) > self._max_size_bytes:
            raise DocumentTooLargeError(
                f"Document exceeds the {self._max_size_bytes // (1024 * 1024)} MB limit"
            )

        try:
            text = self._extract_text(suffix, content).strip()
        except Exception as exc:
            raise InvalidDocumentError("The uploaded document could not be read") from exc
        if not text:
            raise EmptyDocumentError("No readable text was found in the document")

        user_directory = self._upload_dir / str(user_id)
        user_directory.mkdir(parents=True, exist_ok=True)
        stored_path = user_directory / f"{uuid4().hex}_{safe_name}"
        stored_path.write_bytes(content)
        return StoredDocument(original_name=safe_name, path=stored_path, text=text)

    @staticmethod
    def _safe_file_name(file_name: str) -> str:
        base_name = Path(file_name or "document").name
        normalized = re.sub(r"[^A-Za-z0-9._-]+", "_", base_name).strip("._")
        return normalized or "document"

    @staticmethod
    def _extract_text(suffix: str, content: bytes) -> str:
        if suffix == ".pdf":
            reader = PdfReader(BytesIO(content))
            return "\n".join(page.extract_text() or "" for page in reader.pages)

        document = Document(BytesIO(content))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        table_rows = [
            " | ".join(cell.text for cell in row.cells)
            for table in document.tables
            for row in table.rows
        ]
        return "\n".join([*paragraphs, *table_rows])
